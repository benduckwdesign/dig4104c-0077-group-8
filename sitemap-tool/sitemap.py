import asyncio
from time import sleep
from pyppeteer import launch
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement

links = []

@dataclass()
class NavigationItem:
    name: field(init=True, default_factory=str)
    link: field(init=True, default_factory=str)
    parent: field(init=True, default_factory=str)
    sherlock: field(init=True, default_factory=int)
    paragraph: field(init=True, default_factory=None)
    # children: field(default_factory=list)

    def print_description(self):
        return """=== Link ===
    name: %s
    link: %s
    parent: %s
--- --- ---""" % (self.name, self.link, self.parent)

    def __eq__(self, other):
        return self.name == other.name and self.link == other.link

async def upsertInLinksDatabase(link_obj):
    global links
    match = 0
    for curr_link_obj in links:
        if link_obj == curr_link_obj:
            curr_link_obj.sherlock = curr_link_obj.sherlock + 1
            match = 1
            break
    if match == 0:
        links.append(link_obj)
        print(link_obj.print_description())

async def scanForLinks(page):
    global links

    # Get sublink menu parent (if there is one)
    parent_name = ""
    try:
        parent_menu = await page.querySelector('.ptnav2selected')
        parent_text = await parent_menu.getProperty('textContent')
        parent_name = await parent_text.jsonValue()
    except:
        parent_name = ""

    # Refresh and parse anchors
    all_anchors = await page.querySelectorAll('a')
    for anchor in all_anchors:
        anchor_href = await anchor.getProperty('href')
        anchor_text = await anchor.getProperty('textContent')
        href_value = await anchor_href.jsonValue()
        text_value = await anchor_text.jsonValue()
        if href_value != "" and text_value != "":
            if parent_name == text_value:
                await upsertInLinksDatabase(NavigationItem(text_value, href_value, "", 0, None))
            else:
                await upsertInLinksDatabase(NavigationItem(text_value, href_value, parent_name, 0, None))
            #links.add(""+text_value+" => "+href_value)
            #i luv queen emeraldas

async def expandToggles(page, exec):
    # Expand all sections (only works one by one...)
    nav_toggles = await page.querySelectorAll('.ptnav2toggle')
    for toggle in nav_toggles:
        await toggle.hover()
        await toggle.click()

        await page.waitFor(1000)

        #await page.screenshot({'path': 'page1.png', 'fullPage': True})

        await exec(page)

        nav_toggles2 = await page.querySelectorAll('.ptnav2toggle')
        for toggle2 in nav_toggles2:
                await toggle2.hover()
                await toggle2.click()

                await page.waitFor(1000)

                #await page.screenshot({'path': 'page1.png', 'fullPage': True})

                await exec(page)

# def insert_paragraph_after(paragraph, text=None, style=None):
#     """Insert a new paragraph after the given paragraph."""
#     new_p = OxmlElement("w:p")
#     paragraph._p.addnext(new_p)
#     new_para = Paragraph(new_p, paragraph._parent)
#     if text:
#         new_para.add_run(text)
#     if style is not None:
#         new_para.style = style
#     return new_para

async def loginToMyUCF(browser):
    page = await browser.newPage()

    await page.setUserAgent("Mozilla/5.0 (Windows NT 10.0; Win64; x64)")

    print("Logging in.")
    await page.goto('https://my.ucf.edu/')

    await page.click('form button')
    #await page.waitForNavigation()

    #await page.goto('https://idp-prod.cc.ucf.edu/idp/profile/SAML2/Redirect/SSO?execution=e3s1')

    #await page.focus('#username')

    await page.waitForSelector('#username')

    try:
        with open("./username","r") as file:
            file.seek(0)
            username = file.readline()
            await page.type('#username', username)
    except:
        print("Please put your username in a text file named 'username' (no extension.)")

    try:
        with open("./password","r") as file:
            file.seek(0)
            password = file.readline()
            await page.type('#password', password)
    except:
        print("Please put your password in a text file named 'password' (no extension.)")

    await page.click('.btn')

    print("Waiting for homepage to load.")

    await page.waitForNavigation({ 'waitUntil': 'networkidle0'})

    await page.waitForSelector('#fldra_FX_STUDENT_SLFSRV_MENU_90')

    return page

async def exportLinksToDocument(links, title):
    document = Document()
    document.add_heading(title, 0)

    def sort_link_by_sherlock(link):
        return link.sherlock

    sorted_links = sorted(links, key=sort_link_by_sherlock, reverse=True)

    for link in sorted_links:

        if link.parent == "":
            link.paragraph = document.add_paragraph(
                "%s" % (link.name), style='List Bullet'
            )
            link.paragraph.paragraph_format.left_indent = Inches(0)
    
    for link in sorted_links:

        if link.parent != "":
            for parent_link in sorted_links:
                if link.parent == parent_link.name:
                    link.paragraph = document.add_paragraph(
                        "%s" % (link.name), style='List Bullet'
                    )
                    link.paragraph.paragraph_format.left_indent = parent_link.paragraph.paragraph_format.left_indent + Inches(0.5)
                    parent_link.paragraph._p.addnext(link.paragraph._p)

    document.save('%s.docx' % (title))

async def exportTextToDocument(text_lines, title):
    document = Document()
    document.add_heading(title, 0)
    for line in text_lines:
        p = document.add_paragraph(
            "%s" % (line), style="List Bullet"
        )
        p.paragraph_format.left_indent = Inches(0)
    document.save('%s.docx' % (title))

async def main():
    global links

    print("Launching headless browser.")
    browser = await launch()

    page = await loginToMyUCF(browser)

    print("Scanning links. Please wait. This will take a while.")

    await expandToggles(page, scanForLinks)

    print("Found %s unique links." % (len(links)))

    #print("Searching ...") #TODO click on and scan additional subpages for student center
    
    #await page.goto('https://my.ucf.edu/psp/IHPROD/EMPLOYEE/CSPROD/c/SA_LEARNER_SERVICES.SSS_STUDENT_CENTER.GBL?pt_fname=FX_STUDENT_SLFSRV_MENU_90&FolderPath=PORTAL_ROOT_OBJECT.FX_STUDENT_SLFSRV_MENU_90&IsFolder=true')
    await page.waitForSelector('a#fldra_FX_STUDENT_SLFSRV_MENU_90')
    studentselfsrv = await page.querySelector('a#fldra_FX_STUDENT_SLFSRV_MENU_90')
    await studentselfsrv.click()
    #await page.waitForSelector('#DERIVED_SSS_SCR_SSS_LINK_ANCHOR4')
    await page.waitForNavigation({ 'waitUntil': 'networkidle0', 'timeout': '0'})

    student_center_links = await page.querySelectorAll('a.PSHYPERLINK')
    student_center_link_names = []
    for link in student_center_links:
        link_text = await link.getProperty('textContent')
        link_name = await link_text.jsonValue()
        student_center_link_names.append(link_name)
        print(link_name)
    
    await page.screenshot({'path': "./screenshot.png", 'fullPage': 'true'})

    print("Closing browser.")
    await browser.close()

    print("Exporting links to Word document file.")
    await exportLinksToDocument(links, "myUCF_Homepage")
    await exportTextToDocument(student_center_link_names, "myUCF_Student_Center_Links")
    

asyncio.get_event_loop().run_until_complete(main())